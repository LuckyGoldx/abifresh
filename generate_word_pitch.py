"""
Generate comprehensive Hospital EMR Pitch Document in Word format.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "Hospital_EMR_Pitch")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

# ─── Styles ───────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 51, 102)  # dark blue

def add_colored_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_flowchart_text(doc, steps, title=""):
    """Add a simple text-based flowchart."""
    if title:
        doc.add_paragraph(title, style='Heading 3')
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"┌─────────────────────────────────┐")
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"│  {step:<33s}│")
        run2.font.name = 'Consolas'
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0, 51, 102)
        
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"└─────────────────────────────────┘")
        run3.font.name = 'Consolas'
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0, 51, 102)
        
        if i < len(steps) - 1:
            arrow = doc.add_paragraph()
            arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            a = arrow.add_run("▼")
            a.font.size = Pt(14)
            a.font.color.rgb = RGBColor(0, 102, 153)

# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("🏥")
run.font.size = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HOSPITAL EMR SYSTEM")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Electronic Medical Records & Hospital Management Platform")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Comprehensive Solution for Private Hospital Operations")
run.font.size = Pt(14)
run.italic = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL PITCH DOCUMENT")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(153, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared: February 2026  |  Version 1.0")
run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS 
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('TABLE OF CONTENTS', level=1)
toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "System Overview & Architecture"),
    ("3.", "Core Modules & Features"),
    ("  3.1", "Patient Management"),
    ("  3.2", "Electronic Medical Records (EMR)"),
    ("  3.3", "Appointment & Scheduling"),
    ("  3.4", "Staff & HR Management"),
    ("  3.5", "Billing & Financial Management"),
    ("  3.6", "Pharmacy Management"),
    ("  3.7", "Laboratory Management"),
    ("  3.8", "Radiology & Imaging"),
    ("  3.9", "Inventory & Supply Chain"),
    ("  3.10", "Nursing Station"),
    ("  3.11", "In-Patient (Ward) Management"),
    ("  3.12", "Out-Patient Department (OPD)"),
    ("  3.13", "Emergency Department"),
    ("  3.14", "Operating Theatre Management"),
    ("  3.15", "Blood Bank Management"),
    ("4.", "Administrative Modules"),
    ("  4.1", "Admin Dashboard & Analytics"),
    ("  4.2", "Reports & Business Intelligence"),
    ("  4.3", "User & Role Management"),
    ("  4.4", "Audit Trail & Compliance"),
    ("5.", "Workflow Diagrams"),
    ("6.", "Integration & Interoperability"),
    ("7.", "Security & Compliance"),
    ("8.", "Technology Stack"),
    ("9.", "Implementation Plan"),
    ("10.", "Pricing & Licensing"),
    ("11.", "Support & Maintenance"),
    ("12.", "Why Choose Our EMR?"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(11)
    if not num.startswith(" "):
        run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
doc.add_paragraph(
    "Our Hospital Electronic Medical Records (EMR) system is a comprehensive, integrated "
    "healthcare management platform designed specifically for private hospitals. It digitizes "
    "every aspect of hospital operations — from patient registration to discharge, from pharmacy "
    "dispensing to laboratory results, from staff scheduling to financial reporting."
)
doc.add_paragraph(
    "The system eliminates paper-based processes, reduces medical errors, improves patient outcomes, "
    "and provides real-time visibility into hospital performance. Built with modern technologies, "
    "it is scalable, secure, and compliant with healthcare data protection standards."
)

doc.add_heading('Key Value Propositions', level=2)
benefits = [
    ("✅ Paperless Operations", "Eliminate paper records, forms, and manual tracking entirely"),
    ("✅ Improved Patient Safety", "Drug interaction alerts, allergy warnings, clinical decision support"),
    ("✅ Revenue Optimization", "Automated billing, insurance claims, and financial analytics"),
    ("✅ Operational Efficiency", "Streamlined workflows reduce wait times by up to 60%"),
    ("✅ Regulatory Compliance", "Full audit trails, data encryption, role-based access control"),
    ("✅ Data-Driven Decisions", "Real-time dashboards, KPIs, and business intelligence reports"),
    ("✅ Staff Productivity", "Automated scheduling, task assignment, and workload balancing"),
    ("✅ Patient Satisfaction", "Online booking, patient portal, SMS/email notifications"),
]
add_colored_table(doc, ["Benefit", "Description"], benefits)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. SYSTEM OVERVIEW & ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('2. SYSTEM OVERVIEW & ARCHITECTURE', level=1)

doc.add_heading('High-Level Architecture', level=2)
doc.add_paragraph(
    "The Hospital EMR system follows a modern three-tier architecture with clear separation "
    "of concerns, ensuring scalability, maintainability, and security."
)

# Architecture diagram as text
arch_lines = [
    "╔══════════════════════════════════════════════════════════════════════╗",
    "║                    HOSPITAL EMR SYSTEM ARCHITECTURE                 ║",
    "╠══════════════════════════════════════════════════════════════════════╣",
    "║                                                                    ║",
    "║  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐             ║",
    "║  │  Web Browser  │  │ Mobile App   │  │  Tablet/Kiosk│             ║",
    "║  │  (Desktop)    │  │ (iOS/Android)│  │  (Check-in)  │             ║",
    "║  └──────┬───────┘  └──────┬───────┘  └──────┬───────┘             ║",
    "║         │                 │                  │                     ║",
    "║         └────────────────┬┴─────────────────┘                     ║",
    "║                          │                                        ║",
    "║              ┌───────────▼───────────┐                            ║",
    "║              │   API GATEWAY / LB     │                           ║",
    "║              │   (NGINX / Traefik)    │                           ║",
    "║              └───────────┬───────────┘                            ║",
    "║                          │                                        ║",
    "║    ┌─────────────────────┼─────────────────────┐                  ║",
    "║    │                     │                     │                  ║",
    "║    ▼                     ▼                     ▼                  ║",
    "║  ┌──────────┐  ┌──────────────┐  ┌──────────────┐               ║",
    "║  │Auth Svc  │  │ Core EMR API │  │ Analytics Svc│               ║",
    "║  │(JWT/RBAC)│  │ (REST/GQL)   │  │ (Reports)   │               ║",
    "║  └────┬─────┘  └──────┬───────┘  └──────┬──────┘               ║",
    "║       │               │                  │                      ║",
    "║       └───────────────┼──────────────────┘                      ║",
    "║                       │                                         ║",
    "║              ┌────────▼────────┐                                ║",
    "║              │   DATABASE      │                                ║",
    "║              │  (PostgreSQL)   │                                ║",
    "║              │  + Redis Cache  │                                ║",
    "║              │  + File Storage │                                ║",
    "║              └─────────────────┘                                ║",
    "║                                                                 ║",
    "╚═════════════════════════════════════════════════════════════════════╝",
]
for line in arch_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph()
doc.add_heading('System Components', level=2)
components = [
    ("Frontend (Presentation)", "React.js / Next.js", "Web dashboard, patient portal, admin panels"),
    ("Mobile Application", "React Native / Flutter", "Doctor rounds, nurse stations, patient app"),
    ("API Layer", "Node.js / Python FastAPI", "RESTful APIs, GraphQL, WebSocket real-time"),
    ("Database", "PostgreSQL + Redis", "Relational data, caching, session management"),
    ("File Storage", "MinIO / AWS S3", "Medical images, documents, reports, X-rays"),
    ("Authentication", "JWT + OAuth 2.0", "SSO, MFA, role-based access control"),
    ("Message Queue", "RabbitMQ / Redis Pub/Sub", "Async tasks, notifications, lab results"),
    ("Monitoring", "Grafana + Prometheus", "System health, performance metrics, alerts"),
]
add_colored_table(doc, ["Component", "Technology", "Purpose"], components)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. CORE MODULES & FEATURES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('3. CORE MODULES & FEATURES', level=1)

# Module overview diagram
doc.add_paragraph(
    "The Hospital EMR system consists of 15+ integrated modules covering every "
    "department and function of a private hospital."
)

module_overview = [
    "╔════════════════════════════════════════════════════════════════╗",
    "║              HOSPITAL EMR - MODULE OVERVIEW                   ║",
    "╠════════════════════════════════════════════════════════════════╣",
    "║                                                               ║",
    "║  CLINICAL MODULES          ADMINISTRATIVE MODULES             ║",
    "║  ┌─────────────────┐       ┌─────────────────┐               ║",
    "║  │ Patient Mgmt    │       │ Admin Dashboard  │               ║",
    "║  │ EMR / EHR       │       │ User Management  │               ║",
    "║  │ OPD Management  │       │ Reports / BI     │               ║",
    "║  │ IPD / Wards     │       │ Audit Trail      │               ║",
    "║  │ Emergency Dept  │       │ Settings/Config  │               ║",
    "║  │ Operating Theatre│      └─────────────────┘               ║",
    "║  │ Nursing Station │                                          ║",
    "║  └─────────────────┘       FINANCIAL MODULES                  ║",
    "║                            ┌─────────────────┐               ║",
    "║  SUPPORT MODULES           │ Billing/Invoicing│               ║",
    "║  ┌─────────────────┐       │ Insurance Claims │               ║",
    "║  │ Pharmacy        │       │ Payroll          │               ║",
    "║  │ Laboratory      │       │ Accounting       │               ║",
    "║  │ Radiology       │       └─────────────────┘               ║",
    "║  │ Blood Bank      │                                          ║",
    "║  │ Inventory       │       HR & STAFF MODULES                 ║",
    "║  │ Diet/Kitchen    │       ┌─────────────────┐               ║",
    "║  └─────────────────┘       │ Staff Management │               ║",
    "║                            │ Scheduling/Roster│               ║",
    "║                            │ Attendance       │               ║",
    "║                            │ Performance      │               ║",
    "║                            └─────────────────┘               ║",
    "╚════════════════════════════════════════════════════════════════╝",
]
for line in module_overview:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph()
doc.add_page_break()

# ─── 3.1 Patient Management ──────────────────────────────────────────
doc.add_heading('3.1 Patient Management', level=2)
doc.add_paragraph(
    "Complete patient lifecycle management from first contact through discharge and follow-up. "
    "Central to the entire EMR system, providing a unified patient profile accessible across all departments."
)

patient_features = [
    ("Patient Registration", "Demographics, contact info, emergency contacts, photo capture, biometric ID"),
    ("Unique Patient ID (UHID)", "Auto-generated unique hospital ID with barcode/QR code"),
    ("Patient Search", "Quick search by name, ID, phone, national ID — instant results"),
    ("Patient Categories", "VIP, regular, insurance, corporate, staff, pediatric, geriatric"),
    ("Medical History", "Past illnesses, surgeries, family history, social history, allergies"),
    ("Insurance Management", "Multiple insurance policies, pre-authorization, coverage limits"),
    ("Document Upload", "ID cards, insurance cards, referral letters, consent forms"),
    ("Patient Portal", "Self-service: view records, book appointments, download reports"),
    ("Patient Merge/Link", "Merge duplicate records, link family members"),
    ("Communication", "SMS, email, WhatsApp notifications for appointments and results"),
]
add_colored_table(doc, ["Feature", "Description"], patient_features)

doc.add_paragraph()
add_flowchart_text(doc, [
    "Patient Arrives at Hospital",
    "Registration Desk / Self-Kiosk",
    "Search Existing Records",
    "New? Create Profile + UHID",
    "Capture Demographics & Insurance",
    "Issue Patient Card (QR/Barcode)",
    "Direct to Appropriate Department",
], "Patient Registration Workflow")

doc.add_page_break()

# ─── 3.2 Electronic Medical Records ──────────────────────────────────
doc.add_heading('3.2 Electronic Medical Records (EMR)', level=2)
doc.add_paragraph(
    "The heart of the system — a complete digital health record for every patient. "
    "Enables clinicians to document, review, and share patient information securely."
)

emr_features = [
    ("Clinical Notes", "SOAP notes, progress notes, consultation notes with templates"),
    ("Diagnosis (ICD-10/11)", "Searchable ICD coding, multiple diagnoses per encounter"),
    ("Prescriptions", "E-prescribing with drug interaction checks, dosage calculators"),
    ("Vitals Recording", "BP, temperature, pulse, SpO2, weight, height, BMI auto-calc"),
    ("Order Management", "Lab orders, radiology orders, procedure orders — all digital"),
    ("Clinical Templates", "Specialty-specific templates (cardiology, orthopedics, etc.)"),
    ("Allergy & Alert System", "Drug allergies, food allergies, condition alerts — prominent display"),
    ("Document Generation", "Auto-generate discharge summaries, referral letters, certificates"),
    ("Medical Timeline", "Chronological view of all patient encounters, results, procedures"),
    ("Clinical Decision Support", "Evidence-based alerts, guideline reminders, risk scores"),
    ("Voice-to-Text", "Dictation support for clinical notes (with AI transcription)"),
    ("E-Signatures", "Digital signatures for prescriptions, consent forms, notes"),
]
add_colored_table(doc, ["Feature", "Description"], emr_features)

doc.add_paragraph()
add_flowchart_text(doc, [
    "Patient Encounter Begins",
    "Pull Up Patient Record (UHID)",
    "Review History & Alerts",
    "Record Vitals & Complaints",
    "Clinical Examination & SOAP Note",
    "Order Labs / Imaging / Procedures",
    "Prescribe Medications",
    "Generate Encounter Summary",
    "Schedule Follow-up / Discharge",
], "Clinical Documentation Workflow")

doc.add_page_break()

# ─── 3.3 Appointment & Scheduling ────────────────────────────────────
doc.add_heading('3.3 Appointment & Scheduling', level=2)
doc.add_paragraph(
    "Intelligent appointment scheduling system that manages doctor availability, "
    "reduces no-shows, and optimizes resource utilization."
)

appt_features = [
    ("Online Booking", "Patients book via web portal or mobile app — 24/7"),
    ("Walk-in Queue", "Real-time queue management with estimated wait times"),
    ("Doctor Schedule", "Define availability slots, breaks, leave, on-call schedules"),
    ("Multi-Resource Booking", "Book doctor + room + equipment simultaneously"),
    ("Recurring Appointments", "Dialysis, physiotherapy — auto-schedule recurring visits"),
    ("Reminders & Alerts", "SMS/email/WhatsApp reminders 24h and 1h before appointment"),
    ("No-Show Tracking", "Track no-shows, auto-waitlist, reschedule management"),
    ("Calendar Integration", "Sync with Google Calendar, Outlook for doctors"),
    ("Telemedicine Slots", "Video consultation slots with integrated video calling"),
    ("Capacity Dashboard", "Real-time view of clinic capacity and utilization"),
]
add_colored_table(doc, ["Feature", "Description"], appt_features)

doc.add_paragraph()
add_flowchart_text(doc, [
    "Patient Requests Appointment",
    "Select Doctor / Specialty / Date",
    "Check Availability (Real-time)",
    "Confirm & Book Slot",
    "Send Confirmation (SMS/Email)",
    "Reminder Sent (24h before)",
    "Patient Arrives → Check-in",
    "Queue to Doctor Consultation",
], "Appointment Booking Workflow")

doc.add_page_break()

# ─── 3.4 Staff & HR Management ───────────────────────────────────────
doc.add_heading('3.4 Staff & HR Management', level=2)
doc.add_paragraph(
    "Complete human resource management for all hospital staff — doctors, nurses, "
    "technicians, administrators, and support personnel."
)

hr_features = [
    ("Employee Profiles", "Complete profiles with credentials, certifications, specializations"),
    ("Recruitment Pipeline", "Job postings, applications, interview scheduling, offer management"),
    ("Attendance & Time", "Clock-in/clock-out, biometric integration, overtime tracking"),
    ("Shift Scheduling", "Auto-generate shift rosters, swap requests, on-call management"),
    ("Leave Management", "Apply, approve, track leave — annual, sick, maternity, emergency"),
    ("Payroll Integration", "Salary calculation, deductions, bonuses, pay slip generation"),
    ("Performance Reviews", "360° evaluations, KPIs, competency assessments"),
    ("Training & CME", "Track continuing medical education, certifications, expiry alerts"),
    ("Credentialing", "Verify and track medical licenses, board certifications"),
    ("Department Management", "Organizational hierarchy, department heads, reporting lines"),
    ("Staff Communication", "Internal messaging, announcements, policy distribution"),
    ("Expense Management", "Travel claims, reimbursements, petty cash management"),
]
add_colored_table(doc, ["Feature", "Description"], hr_features)

doc.add_page_break()

# ─── 3.5 Billing & Financial Management ──────────────────────────────
doc.add_heading('3.5 Billing & Financial Management', level=2)
doc.add_paragraph(
    "End-to-end financial management covering patient billing, insurance claims, "
    "revenue cycle management, and accounting integration."
)

billing_features = [
    ("Patient Billing", "Auto-generate bills from orders, procedures, bed charges, consumables"),
    ("Insurance Claims", "Electronic claim submission, pre-authorization, adjudication tracking"),
    ("Multiple Payment Methods", "Cash, card, mobile money, bank transfer, installment plans"),
    ("Price Lists / Tariffs", "Configurable price lists for services, by department or package"),
    ("Package Billing", "Health check packages, surgery packages, maternity packages"),
    ("Deposit Management", "Advance deposits, refunds, deposit adjustment tracking"),
    ("Credit Management", "Corporate accounts, credit limits, outstanding tracking"),
    ("Revenue Dashboard", "Daily/weekly/monthly revenue, department-wise collection"),
    ("Tax Management", "VAT/GST calculation, tax invoices, exemptions"),
    ("Refund Processing", "Structured refund approval workflow, audit trail"),
    ("Financial Reports", "P&L, balance sheet, aging reports, revenue forecasts"),
    ("Accounting Integration", "Export to QuickBooks, Xero, Tally, or custom ERP"),
]
add_colored_table(doc, ["Feature", "Description"], billing_features)

doc.add_paragraph()
add_flowchart_text(doc, [
    "Services Rendered (Consult/Lab/Proc)",
    "Charges Auto-Captured to Bill",
    "Insurance? → Pre-Authorization",
    "Generate Invoice",
    "Patient Payment (Cash/Card/Mobile)",
    "Insurance Claim Submitted",
    "Payment Reconciliation",
    "Receipt Generated",
], "Billing Workflow")

doc.add_page_break()

# ─── 3.6 Pharmacy Management ─────────────────────────────────────────
doc.add_heading('3.6 Pharmacy Management', level=2)
doc.add_paragraph(
    "Integrated pharmacy system for both in-patient and out-patient dispensing "
    "with full inventory tracking and drug safety features."
)

pharmacy_features = [
    ("E-Prescription Receipt", "Receive prescriptions electronically from EMR — no paper"),
    ("Drug Dispensing", "Barcode-based dispensing, patient verification, label printing"),
    ("Drug Interaction Check", "Real-time alerts for drug-drug & drug-allergy interactions"),
    ("Inventory Management", "Stock levels, reorder points, expiry tracking, batch management"),
    ("Formulary Management", "Drug formulary, generic substitution suggestions"),
    ("Controlled Substances", "Special tracking for narcotics/controlled medications"),
    ("Return Management", "Process drug returns, restocking, wastage tracking"),
    ("Purchase Orders", "Auto-generate POs when stock falls below reorder level"),
    ("Supplier Management", "Vendor database, price comparison, delivery tracking"),
    ("Prescription History", "Full dispensing history per patient for safety review"),
    ("Point of Sale", "Walk-in customer sales, OTC medications"),
    ("Reports", "Fast/slow movers, expiry reports, revenue by category"),
]
add_colored_table(doc, ["Feature", "Description"], pharmacy_features)

doc.add_page_break()

# ─── 3.7 Laboratory Management ───────────────────────────────────────
doc.add_heading('3.7 Laboratory Management (LIS)', level=2)
doc.add_paragraph(
    "Complete Laboratory Information System for sample collection, processing, "
    "result entry, and reporting — integrated with the EMR."
)

lab_features = [
    ("Order Management", "Receive lab orders digitally from doctors, priority flagging"),
    ("Sample Collection", "Barcode labels, collection tracking, phlebotomy scheduling"),
    ("Sample Tracking", "Track sample from collection → processing → reporting"),
    ("Result Entry", "Manual or instrument-interfaced result capture with normal ranges"),
    ("Auto-Validation", "Rules-based auto-validation for results within normal ranges"),
    ("Critical Alerts", "Immediate alerts for critical/panic values to ordering doctor"),
    ("Report Generation", "Professional lab reports with reference ranges, interpretations"),
    ("Test Catalog", "Comprehensive test catalog with panels, profiles, and pricing"),
    ("Quality Control", "QC sample tracking, Levey-Jennings charts, Westgard rules"),
    ("Instrument Integration", "Bi-directional interface with lab analyzers"),
    ("TAT Monitoring", "Track turnaround times, identify bottlenecks"),
    ("External Lab Referral", "Send samples to reference labs, track results"),
]
add_colored_table(doc, ["Feature", "Description"], lab_features)

doc.add_paragraph()
add_flowchart_text(doc, [
    "Doctor Orders Lab Test",
    "Order Appears in Lab Queue",
    "Sample Collection (Barcode Label)",
    "Sample Processing (Analyzer)",
    "Result Entry / Auto-Capture",
    "Validation (Auto + Manual QC)",
    "Report Generated → EMR Updated",
    "Doctor & Patient Notified",
], "Laboratory Workflow")

doc.add_page_break()

# ─── 3.8 Radiology & Imaging ─────────────────────────────────────────
doc.add_heading('3.8 Radiology & Imaging (RIS/PACS)', level=2)
doc.add_paragraph(
    "Radiology Information System with Picture Archiving and Communication System "
    "for managing imaging orders, images, and reports."
)

radiology_features = [
    ("Order Management", "Digital radiology orders with clinical indication"),
    ("Scheduling", "Schedule imaging appointments, manage equipment slots"),
    ("Worklist (DICOM MWL)", "DICOM Modality Worklist integration with imaging equipment"),
    ("Image Storage (PACS)", "DICOM image storage, viewing, and archiving"),
    ("Image Viewer", "Web-based DICOM viewer with measurement tools"),
    ("Report Dictation", "Structured reporting templates, voice dictation support"),
    ("Report Distribution", "Auto-send reports to EMR, referring doctor, patient portal"),
    ("Comparison Studies", "Side-by-side comparison of current and prior images"),
    ("CD/USB Burning", "Export images to CD/USB for patient or referral"),
    ("Equipment Tracking", "Maintenance schedules, utilization reports, downtime tracking"),
]
add_colored_table(doc, ["Feature", "Description"], radiology_features)

doc.add_page_break()

# ─── 3.9 Inventory & Supply Chain ────────────────────────────────────
doc.add_heading('3.9 Inventory & Supply Chain Management', level=2)
doc.add_paragraph(
    "Centralized inventory management for all hospital supplies — from surgical "
    "consumables to office stationery, with automated reordering."
)

inv_features = [
    ("Multi-Store Management", "Central store, department sub-stores, pharmacy, kitchen"),
    ("Item Master", "Comprehensive catalog with categories, units, specifications"),
    ("Stock Management", "Real-time stock levels, bin location tracking"),
    ("Purchase Requisition", "Department-wise requisitions with approval workflows"),
    ("Purchase Orders", "Vendor selection, PO generation, delivery tracking"),
    ("Goods Receipt", "GRN with quality inspection, batch/lot tracking"),
    ("Issue/Transfer", "Inter-store transfers, department issue tracking"),
    ("Expiry Management", "FEFO/FIFO tracking, expiry alerts 30/60/90 days"),
    ("ABC/VED Analysis", "Classify items by value and criticality for optimization"),
    ("Barcode/RFID", "Item identification and tracking via barcode or RFID"),
    ("Vendor Management", "Vendor evaluation, price comparison, contract management"),
    ("Reports", "Consumption reports, stock valuation, reorder analysis"),
]
add_colored_table(doc, ["Feature", "Description"], inv_features)

doc.add_page_break()

# ─── 3.10 Nursing Station ────────────────────────────────────────────
doc.add_heading('3.10 Nursing Station', level=2)
doc.add_paragraph(
    "Dedicated nursing workflow tools for patient care documentation, "
    "medication administration, and task management."
)

nursing_features = [
    ("Patient Worklist", "View assigned patients with alerts and pending tasks"),
    ("Nursing Assessment", "Admission assessment, pain scales, fall risk, pressure ulcer risk"),
    ("Vital Signs Charting", "Record and chart vitals with auto-alerts for abnormals"),
    ("Medication Administration", "eMAR — barcode scan patient + medication for safety"),
    ("Care Plans", "Create and manage individualized care plans"),
    ("Intake/Output", "Track fluid intake, urine output, drain output"),
    ("Wound Care", "Document wound assessments with photos, treatment tracking"),
    ("Shift Handover", "Structured handover notes with critical information flagging"),
    ("Doctor Rounding", "Assist with ward rounds — display patient summary"),
    ("Task Management", "Auto-assigned tasks from orders, with completion tracking"),
]
add_colored_table(doc, ["Feature", "Description"], nursing_features)

doc.add_page_break()

# ─── 3.11 In-Patient (Ward) Management ───────────────────────────────
doc.add_heading('3.11 In-Patient (Ward) Management', level=2)
doc.add_paragraph(
    "Complete in-patient management from admission through discharge, "
    "covering bed management, ward operations, and discharge planning."
)

ipd_features = [
    ("Admission", "Direct, emergency, or transfer admission with bed assignment"),
    ("Bed Management", "Visual bed map, occupancy dashboard, bed categories (ICU, general, private)"),
    ("Ward Dashboard", "Real-time ward status, patient list, alerts"),
    ("Transfer Management", "Ward-to-ward, bed-to-bed transfers with charge adjustments"),
    ("Diet Management", "Dietary orders, special diets, kitchen integration"),
    ("Discharge Planning", "Discharge summary, follow-up appointments, take-home meds"),
    ("Discharge Checklist", "Ensure all clearances (billing, pharmacy, lab) before discharge"),
    ("Bed Charges", "Automatic daily bed charge calculation based on bed type"),
    ("Visitor Management", "Track visitors, visiting hours, visitor passes"),
    ("Census Reports", "Daily census, admission/discharge/transfer (ADT) reports"),
]
add_colored_table(doc, ["Feature", "Description"], ipd_features)

doc.add_paragraph()
add_flowchart_text(doc, [
    "Admission Request (OPD/ER/Direct)",
    "Check Bed Availability",
    "Assign Bed (Ward/Private/ICU)",
    "Admission Documentation",
    "Daily Rounds & Orders",
    "Nursing Care & Monitoring",
    "Discharge Planning Initiated",
    "Discharge Clearance (All Depts)",
    "Generate Discharge Summary",
    "Patient Discharged + Follow-up",
], "In-Patient Workflow")

doc.add_page_break()

# ─── 3.12 Out-Patient Department ─────────────────────────────────────
doc.add_heading('3.12 Out-Patient Department (OPD)', level=2)
doc.add_paragraph(
    "Streamlined outpatient workflow from arrival to consultation to checkout, "
    "designed to minimize patient wait times."
)

opd_features = [
    ("Check-in / Queue", "Self-service kiosk or receptionist check-in, token system"),
    ("Triage", "Nurse triage with severity classification"),
    ("Doctor Queue", "Display queue for each doctor, estimated wait times"),
    ("Consultation", "Full EMR during consultation — notes, orders, prescriptions"),
    ("Referrals", "Internal referrals to specialists, external referral letters"),
    ("Follow-up Scheduling", "Book follow-up before patient leaves"),
    ("OPD Billing", "Quick billing at checkout — consultation + services"),
    ("Token Display", "Public display screen showing current token numbers"),
    ("Telemedicine", "Virtual OPD consultations via video call"),
    ("OPD Analytics", "Patient footfall, wait times, revenue per doctor"),
]
add_colored_table(doc, ["Feature", "Description"], opd_features)

doc.add_page_break()

# ─── 3.13 Emergency Department ───────────────────────────────────────
doc.add_heading('3.13 Emergency Department', level=2)
doc.add_paragraph(
    "Fast-paced emergency module designed for rapid triage, treatment, "
    "and disposition — with critical alerts and streamlined documentation."
)

er_features = [
    ("Rapid Registration", "Minimal info required — full details captured later"),
    ("Triage (ESI / MTS)", "Emergency Severity Index or Manchester Triage System"),
    ("Trauma Tracking", "Mechanism of injury, body diagram, GCS score"),
    ("Real-time Bed Board", "Visual display of all ER bays/beds with status"),
    ("STAT Orders", "Priority lab and imaging orders with STAT flags"),
    ("Resuscitation Timer", "Timer for CPR, drug administration intervals"),
    ("Disposition", "Admit, discharge, transfer, or refer — one-click"),
    ("Ambulance Tracking", "Track incoming ambulances, ETA, pre-hospital info"),
    ("Mass Casualty Mode", "Activate for disasters — simplified documentation"),
    ("ER Dashboard", "Wait times, occupancy, acuity distribution, LWBS rate"),
]
add_colored_table(doc, ["Feature", "Description"], er_features)

doc.add_page_break()

# ─── 3.14 Operating Theatre Management ───────────────────────────────
doc.add_heading('3.14 Operating Theatre (OT) Management', level=2)
doc.add_paragraph(
    "Surgical workflow management from scheduling to post-operative care, "
    "including resource management and safety checklists."
)

ot_features = [
    ("Surgery Scheduling", "Book OT slots, manage conflicts, priority scheduling"),
    ("Pre-Op Assessment", "Anesthesia assessment, consent forms, preparation checklist"),
    ("WHO Surgical Checklist", "Digital WHO safety checklist — sign-in, time-out, sign-out"),
    ("OT Documentation", "Surgical notes, anesthesia record, implant tracking"),
    ("Resource Management", "Equipment, staff, consumable allocation per surgery"),
    ("Post-Op Monitoring", "Recovery room monitoring, PACU handover"),
    ("Surgeon Preference Cards", "Pre-configured instrument and supply lists per surgeon"),
    ("OT Utilization", "Track OT usage, delays, cancellations, efficiency metrics"),
    ("Implant Registry", "Track implants used, lot numbers, patient traceability"),
    ("Consent Management", "Digital surgical consent with patient signature"),
]
add_colored_table(doc, ["Feature", "Description"], ot_features)

doc.add_page_break()

# ─── 3.15 Blood Bank Management ──────────────────────────────────────
doc.add_heading('3.15 Blood Bank Management', level=2)
doc.add_paragraph(
    "Complete blood bank operations from donor management to cross-matching "
    "to transfusion tracking with full regulatory compliance."
)

blood_features = [
    ("Donor Management", "Donor registration, screening, donation history"),
    ("Blood Collection", "Donation processing, component separation, storage"),
    ("Inventory", "Real-time stock by blood group, component, expiry"),
    ("Cross-Matching", "Cross-match requests, compatibility testing, results"),
    ("Issue / Transfusion", "Issue blood units, track transfusions, reaction monitoring"),
    ("Testing", "Mandatory screening tests (HIV, HBV, HCV, Syphilis, Malaria)"),
    ("Discard Management", "Track and document discarded units with reasons"),
    ("Blood Requests", "Emergency and routine request workflows"),
    ("Reports", "Collection, usage, discard, expiry reports for regulators"),
    ("Camp Management", "Organize and manage blood donation camps"),
]
add_colored_table(doc, ["Feature", "Description"], blood_features)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. ADMINISTRATIVE MODULES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('4. ADMINISTRATIVE MODULES', level=1)

# ─── 4.1 Admin Dashboard & Analytics ─────────────────────────────────
doc.add_heading('4.1 Admin Dashboard & Analytics', level=2)
doc.add_paragraph(
    "Real-time executive dashboard providing a bird's eye view of hospital "
    "operations, KPIs, and actionable insights."
)

admin_features = [
    ("Executive Dashboard", "Hospital-wide KPIs: revenue, patient count, bed occupancy, ER wait"),
    ("Department Dashboards", "Department-specific metrics and performance indicators"),
    ("Financial Overview", "Revenue, expenses, outstanding receivables, cash flow"),
    ("Occupancy Metrics", "Bed occupancy rate, average length of stay, turnover rate"),
    ("Staff Overview", "Staff on duty, leave calendar, workload distribution"),
    ("Patient Statistics", "New vs returning, by department, by diagnosis, by payer"),
    ("Quality Indicators", "Infection rates, readmission rates, mortality rates"),
    ("Custom Widgets", "Drag-and-drop dashboard customization"),
    ("Alert Center", "Critical alerts: stock-outs, equipment downtime, pending approvals"),
    ("Comparative Analysis", "Compare metrics across time periods, departments, doctors"),
]
add_colored_table(doc, ["Feature", "Description"], admin_features)

doc.add_page_break()

# ─── 4.2 Reports & Business Intelligence ─────────────────────────────
doc.add_heading('4.2 Reports & Business Intelligence', level=2)

reports_list = [
    ("Clinical Reports", "Patient census, diagnosis distribution, treatment outcomes, mortality"),
    ("Financial Reports", "Revenue analysis, billing summary, outstanding, department P&L"),
    ("Operational Reports", "OT utilization, bed occupancy, ER turnaround, appointment no-shows"),
    ("Pharmacy Reports", "Drug consumption, expiry, stock valuation, fast/slow movers"),
    ("Lab Reports", "Test volume, TAT compliance, revenue by test, QC summary"),
    ("HR Reports", "Headcount, attrition, overtime, leave utilization, training compliance"),
    ("Inventory Reports", "Stock levels, purchase history, consumption trends, dead stock"),
    ("Regulatory Reports", "Notifiable diseases, birth/death certificates, infection reports"),
    ("Custom Report Builder", "Drag-and-drop report builder with filters, grouping, charts"),
    ("Scheduled Reports", "Auto-generate and email reports on schedule (daily/weekly/monthly)"),
    ("Export Options", "PDF, Excel, CSV, print — with hospital branding"),
]
add_colored_table(doc, ["Report Category", "Details"], reports_list)

doc.add_page_break()

# ─── 4.3 User & Role Management ──────────────────────────────────────
doc.add_heading('4.3 User & Role Management', level=2)

roles_list = [
    ("Administrator", "Full system access, user management, configuration"),
    ("Doctor", "Patient records, prescriptions, orders, clinical notes"),
    ("Nurse", "Vitals, medication administration, nursing notes, care plans"),
    ("Receptionist", "Patient registration, appointments, billing"),
    ("Pharmacist", "Drug dispensing, inventory, prescription verification"),
    ("Lab Technician", "Sample processing, result entry, QC"),
    ("Radiologist", "Image review, report dictation, worklist management"),
    ("Accountant", "Billing, insurance claims, financial reports"),
    ("HR Manager", "Staff management, payroll, attendance, recruitment"),
    ("Custom Roles", "Create any role with granular permission configuration"),
]
add_colored_table(doc, ["Role", "Access Scope"], roles_list)

doc.add_paragraph()
doc.add_paragraph(
    "The RBAC (Role-Based Access Control) system supports: granular permissions at module/feature level, "
    "IP-based access restrictions, time-based access, multi-factor authentication, "
    "session management, and single sign-on (SSO) integration."
)

doc.add_page_break()

# ─── 4.4 Audit Trail & Compliance ────────────────────────────────────
doc.add_heading('4.4 Audit Trail & Compliance', level=2)

audit_features = [
    ("Complete Audit Trail", "Every action logged: who, what, when, where, before/after values"),
    ("Data Access Logging", "Track who accessed which patient record and when"),
    ("Change History", "Full version history of all clinical documents"),
    ("HIPAA Compliance", "Data encryption, access controls, breach notification protocols"),
    ("GDPR Support", "Right to access, right to erasure, data portability"),
    ("Consent Management", "Track and manage patient consents digitally"),
    ("Data Backup", "Automated backups, disaster recovery, point-in-time restore"),
    ("Regulatory Reporting", "Template-based reports for health authorities"),
]
add_colored_table(doc, ["Feature", "Description"], audit_features)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. WORKFLOW DIAGRAMS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('5. COMPREHENSIVE WORKFLOW DIAGRAMS', level=1)

# --- Overall Patient Journey ---
doc.add_heading('5.1 Complete Patient Journey', level=2)
patient_journey = [
    "╔════════════════════════════════════════════════════════════════════════════╗",
    "║                      COMPLETE PATIENT JOURNEY                             ║",
    "╠════════════════════════════════════════════════════════════════════════════╣",
    "║                                                                           ║",
    "║  ┌──────────┐     ┌──────────┐     ┌──────────────┐                      ║",
    "║  │ ARRIVAL   │────▶│REGISTER  │────▶│ TRIAGE /     │                      ║",
    "║  │ (Walk-in/ │     │(New/     │     │ OPD QUEUE    │                      ║",
    "║  │ Appt/ER)  │     │ Existing)│     │              │                      ║",
    "║  └──────────┘     └──────────┘     └──────┬───────┘                      ║",
    "║                                            │                              ║",
    "║                                            ▼                              ║",
    "║                                   ┌──────────────┐                       ║",
    "║                                   │ CONSULTATION  │                       ║",
    "║                                   │ (Doctor Visit)│                       ║",
    "║                                   └───────┬──────┘                       ║",
    "║                                           │                              ║",
    "║               ┌───────────┬───────────┬───┴───┬──────────┐              ║",
    "║               ▼           ▼           ▼       ▼          ▼              ║",
    "║         ┌──────────┐┌──────────┐┌──────────┐┌────────┐┌────────┐       ║",
    "║         │ LAB TESTS ││ IMAGING  ││ PHARMACY ││ADMIT TO││ REFER  │       ║",
    "║         │           ││          ││          ││  WARD  ││EXTERNAL│       ║",
    "║         └─────┬────┘└────┬─────┘└─────┬────┘└───┬────┘└────────┘       ║",
    "║               │          │             │         │                      ║",
    "║               └──────────┴─────────────┘         │                      ║",
    "║                          │                       │                      ║",
    "║                          ▼                       ▼                      ║",
    "║                  ┌──────────────┐        ┌──────────────┐              ║",
    "║                  │ FOLLOW-UP /  │        │  IN-PATIENT  │              ║",
    "║                  │ BILLING /    │        │  CARE CYCLE  │              ║",
    "║                  │ DISCHARGE    │        │  → DISCHARGE │              ║",
    "║                  └──────────────┘        └──────────────┘              ║",
    "╚════════════════════════════════════════════════════════════════════════════╝",
]
for line in patient_journey:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph()

# --- Billing Revenue Cycle ---
doc.add_heading('5.2 Revenue Cycle Management', level=2)
revenue_cycle = [
    "╔══════════════════════════════════════════════════════════════╗",
    "║               REVENUE CYCLE MANAGEMENT                      ║",
    "╠══════════════════════════════════════════════════════════════╣",
    "║                                                             ║",
    "║  ┌────────────┐    ┌────────────┐    ┌────────────┐        ║",
    "║  │ PATIENT     │───▶│ SERVICE    │───▶│ CHARGE     │        ║",
    "║  │ ENCOUNTER   │    │ DELIVERY   │    │ CAPTURE    │        ║",
    "║  └────────────┘    └────────────┘    └─────┬──────┘        ║",
    "║                                            │               ║",
    "║                                            ▼               ║",
    "║  ┌────────────┐    ┌────────────┐    ┌────────────┐        ║",
    "║  │ PAYMENT    │◀───│ CLAIMS     │◀───│ BILLING /  │        ║",
    "║  │ POSTING    │    │ PROCESS    │    │ CODING     │        ║",
    "║  └─────┬──────┘    └────────────┘    └────────────┘        ║",
    "║        │                                                    ║",
    "║        ▼                                                    ║",
    "║  ┌────────────┐    ┌────────────┐                          ║",
    "║  │ DENIAL     │───▶│ REPORTING  │                          ║",
    "║  │ MANAGEMENT │    │ & ANALYTICS│                          ║",
    "║  └────────────┘    └────────────┘                          ║",
    "╚══════════════════════════════════════════════════════════════╝",
]
for line in revenue_cycle:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 102, 51)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph()

# --- Lab Workflow ---
doc.add_heading('5.3 Laboratory Complete Workflow', level=2)
lab_workflow = [
    "╔══════════════════════════════════════════════════════════════════╗",
    "║                  LABORATORY WORKFLOW                             ║",
    "╠══════════════════════════════════════════════════════════════════╣",
    "║                                                                 ║",
    "║  Doctor Orders  ──▶  Lab Receives  ──▶  Sample Collection       ║",
    "║  (via EMR)          Order              (Barcode Label)          ║",
    "║                                            │                    ║",
    "║                                            ▼                    ║",
    "║  Report Sent    ◀──  Validation    ◀──  Processing              ║",
    "║  to Doctor           (QC Check)         (Analyzer)             ║",
    "║       │                                                         ║",
    "║       ▼                                                         ║",
    "║  Patient Notified ──▶ Available in Portal ──▶ Archived          ║",
    "║                                                                 ║",
    "╚══════════════════════════════════════════════════════════════════╝",
]
for line in lab_workflow:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(102, 0, 102)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph()

# --- Emergency Workflow ---
doc.add_heading('5.4 Emergency Department Workflow', level=2)
er_workflow = [
    "╔═════════════════════════════════════════════════════════════════╗",
    "║                EMERGENCY DEPARTMENT WORKFLOW                    ║",
    "╠═════════════════════════════════════════════════════════════════╣",
    "║                                                                ║",
    "║  Ambulance/    ──▶  Rapid         ──▶  Triage                  ║",
    "║  Walk-in             Registration       (ESI Level 1-5)        ║",
    "║                                            │                   ║",
    "║                          ┌─────────────────┼─────────────┐     ║",
    "║                          ▼                 ▼             ▼     ║",
    "║                     ┌─────────┐     ┌──────────┐  ┌─────────┐ ║",
    "║                     │Resuscit.│     │ Acute    │  │ Urgent  │ ║",
    "║                     │(Level 1)│     │(Level 2-3│  │(Level 4)│ ║",
    "║                     └────┬────┘     └────┬─────┘  └────┬────┘ ║",
    "║                          │               │             │      ║",
    "║                          └───────────────┼─────────────┘      ║",
    "║                                          ▼                    ║",
    "║                               ┌──────────────────┐            ║",
    "║                               │   DISPOSITION     │            ║",
    "║                               │ Admit/Discharge/  │            ║",
    "║                               │ Transfer/Observe  │            ║",
    "║                               └──────────────────┘            ║",
    "╚═════════════════════════════════════════════════════════════════╝",
]
for line in er_workflow:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(153, 0, 0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. INTEGRATION & INTEROPERABILITY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('6. INTEGRATION & INTEROPERABILITY', level=1)

integrations = [
    ("HL7 / FHIR", "Healthcare data exchange standards for interoperability"),
    ("DICOM", "Medical imaging standard for radiology equipment integration"),
    ("ICD-10 / ICD-11", "International classification of diseases for coding"),
    ("SNOMED CT", "Standardized clinical terminology for structured data"),
    ("Lab Instruments", "Bi-directional interface via LIS middleware (ASTM/HL7)"),
    ("Insurance Portals", "Electronic claims, pre-auth, eligibility verification"),
    ("Government Systems", "Public health reporting, birth/death registration"),
    ("Accounting Software", "QuickBooks, Xero, Tally, SAP integration"),
    ("Payment Gateways", "Stripe, PayStack, Flutterwave, bank integrations"),
    ("SMS/Email Gateways", "Twilio, SendGrid, Africa's Talking for communications"),
    ("Video Conferencing", "Zoom, Jitsi for telemedicine consultations"),
    ("Biometric Devices", "Fingerprint readers, facial recognition for attendance"),
    ("Pharmacy Chains", "Integration with external pharmacy networks"),
    ("National Health ID", "Link with national health identification systems"),
]
add_colored_table(doc, ["Integration", "Purpose"], integrations)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. SECURITY & COMPLIANCE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('7. SECURITY & COMPLIANCE', level=1)

security_items = [
    ("Data Encryption", "AES-256 at rest, TLS 1.3 in transit — all data encrypted"),
    ("Access Control", "Role-based (RBAC), attribute-based (ABAC), IP whitelisting"),
    ("Multi-Factor Auth", "SMS OTP, authenticator app, biometric as second factor"),
    ("Session Management", "Auto-timeout, concurrent session limits, device tracking"),
    ("Audit Logging", "Every action logged with timestamp, user, IP, device info"),
    ("Data Backup", "Hourly incremental, daily full, off-site replication"),
    ("Disaster Recovery", "RPO < 1 hour, RTO < 4 hours, automated failover"),
    ("Penetration Testing", "Regular third-party security assessments"),
    ("HIPAA Compliance", "Full compliance with Health Insurance Portability Act"),
    ("GDPR Compliance", "EU data protection — consent, right to erasure, DPO support"),
    ("SOC 2 Type II", "Service organization security controls certification"),
    ("Data Residency", "Choose data storage location for regulatory compliance"),
]
add_colored_table(doc, ["Security Feature", "Details"], security_items)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('8. TECHNOLOGY STACK', level=1)

tech_stack = [
    ("Frontend", "React.js / Next.js 14", "Server-side rendering, responsive UI, PWA support"),
    ("Mobile App", "React Native / Flutter", "Cross-platform iOS & Android, offline-capable"),
    ("Backend API", "Node.js (Express/NestJS) or Python (FastAPI/Django)", "REST + GraphQL APIs"),
    ("Database", "PostgreSQL 16+", "ACID compliant, JSON support, full-text search"),
    ("Caching", "Redis", "Session management, real-time data, message queuing"),
    ("Search Engine", "Elasticsearch", "Full-text search across patient records"),
    ("File Storage", "MinIO / AWS S3", "Medical images, documents, backups"),
    ("Message Queue", "RabbitMQ / Redis Streams", "Async processing, notifications"),
    ("Containerization", "Docker + Kubernetes", "Consistent deployments, auto-scaling"),
    ("CI/CD", "GitHub Actions / GitLab CI", "Automated testing, building, deployment"),
    ("Monitoring", "Prometheus + Grafana", "Metrics, alerting, dashboards"),
    ("Logging", "ELK Stack (Elastic, Logstash, Kibana)", "Centralized log management"),
    ("PACS/DICOM", "Orthanc / dcm4chee", "Open-source DICOM server for imaging"),
]
add_colored_table(doc, ["Layer", "Technology", "Notes"], tech_stack)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. IMPLEMENTATION PLAN
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('9. IMPLEMENTATION PLAN', level=1)
doc.add_paragraph(
    "A phased implementation approach ensures minimal disruption to hospital operations "
    "while delivering value incrementally."
)

phases = [
    ("Phase 1: Foundation\n(Months 1-3)", 
     "• Project setup & infrastructure\n• Patient registration module\n• User management & RBAC\n• Basic EMR (clinical notes)\n• Appointment scheduling\n• Staff profiles"),
    ("Phase 2: Clinical Core\n(Months 4-6)", 
     "• Full EMR with templates\n• OPD workflow\n• Laboratory (LIS)\n• Pharmacy management\n• Billing & invoicing\n• Nursing station"),
    ("Phase 3: Advanced Clinical\n(Months 7-9)", 
     "• In-patient management\n• Emergency department\n• Operating theatre\n• Radiology / PACS\n• Blood bank\n• Diet & kitchen"),
    ("Phase 4: Analytics & Integration\n(Months 10-11)", 
     "• Admin dashboard & KPIs\n• Business intelligence\n• Report builder\n• Insurance integration\n• Payment gateway\n• External system integration"),
    ("Phase 5: Polish & Launch\n(Month 12)", 
     "• Performance optimization\n• Security hardening\n• User training\n• Data migration\n• Go-live support\n• Documentation"),
]
add_colored_table(doc, ["Phase", "Deliverables"], phases)

doc.add_paragraph()

# Implementation timeline diagram
timeline = [
    "╔════════════════════════════════════════════════════════════════════╗",
    "║                   IMPLEMENTATION TIMELINE                         ║",
    "╠════════════════════════════════════════════════════════════════════╣",
    "║                                                                   ║",
    "║  Month:  1  2  3  4  5  6  7  8  9  10 11 12                     ║",
    "║         ├──┼──┼──┼──┼──┼──┼──┼──┼──┼──┼──┤                      ║",
    "║  Phase1 ████████████                                              ║",
    "║  Phase2             ████████████                                  ║",
    "║  Phase3                         ████████████                      ║",
    "║  Phase4                                     ████████              ║",
    "║  Phase5                                             ████          ║",
    "║         ├──┼──┼──┼──┼──┼──┼──┼──┼──┼──┼──┤                      ║",
    "║  Train  ░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░█████████              ║",
    "║  Test   ░░░░░░░████░░░░░████░░░░░████░░░░░████████               ║",
    "║                                                                   ║",
    "╚════════════════════════════════════════════════════════════════════╝",
]
for line in timeline:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. PRICING & LICENSING
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('10. PRICING & LICENSING', level=1)
doc.add_paragraph(
    "Flexible pricing models to suit different hospital sizes and budgets."
)

pricing = [
    ("Starter", "Up to 50 beds", "Core EMR, OPD, Pharmacy, Lab, Billing", "Custom Quote"),
    ("Professional", "50–200 beds", "All Starter + IPD, ER, OT, HR, Analytics", "Custom Quote"),
    ("Enterprise", "200+ beds", "Full suite + PACS, API access, custom dev", "Custom Quote"),
    ("Cloud Hosted", "Any size", "SaaS model — monthly per-user pricing", "Custom Quote"),
    ("On-Premise", "Any size", "Self-hosted with perpetual license", "Custom Quote"),
]
add_colored_table(doc, ["Plan", "Hospital Size", "Includes", "Price"], pricing)

doc.add_paragraph()
doc.add_paragraph(
    "All plans include: implementation support, data migration assistance, "
    "user training (on-site + virtual), 12 months warranty, and technical documentation."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. SUPPORT & MAINTENANCE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('11. SUPPORT & MAINTENANCE', level=1)

support = [
    ("24/7 Help Desk", "Round-the-clock support via phone, email, and chat"),
    ("Dedicated Account Manager", "Single point of contact for your hospital"),
    ("Remote Support", "Screen-sharing for troubleshooting and guidance"),
    ("On-Site Support", "Available for critical issues and installations"),
    ("Regular Updates", "Security patches, feature updates, bug fixes"),
    ("Annual Health Check", "System performance review and optimization"),
    ("Training Sessions", "Quarterly refresher training for staff"),
    ("Knowledge Base", "Comprehensive online documentation and video tutorials"),
    ("SLA Guarantee", "99.9% uptime, <1hr response for critical issues"),
    ("Data Migration", "Expert assistance for migrating from existing systems"),
]
add_colored_table(doc, ["Service", "Description"], support)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 12. WHY CHOOSE OUR EMR?
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('12. WHY CHOOSE OUR EMR?', level=1)

reasons = [
    "✅ Purpose-built for private hospitals — not a generic adaptation",
    "✅ All-in-one solution — no need for multiple vendor systems",
    "✅ Modern technology — fast, responsive, mobile-ready",
    "✅ Scalable — grows with your hospital from 10 beds to 1000+",
    "✅ Interoperable — HL7/FHIR compliant, integrates with anything",
    "✅ Secure — bank-grade encryption, full audit trail",
    "✅ Affordable — flexible licensing, no hidden costs",
    "✅ Local support — dedicated team in your region",
    "✅ Fast implementation — go live in as little as 3 months (Phase 1)",
    "✅ Proven ROI — hospitals report 30-50% efficiency improvement",
]
for reason in reasons:
    p = doc.add_paragraph(reason)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ready to Transform Your Hospital?")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Contact us for a personalized demo and consultation")
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("📧 info@hospitalemr.com  |  📞 +XXX-XXX-XXXX  |  🌐 www.hospitalemr.com")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 102, 153)

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
output_path = os.path.join(OUTPUT_DIR, "Hospital_EMR_Pitch.docx")
doc.save(output_path)
print(f"✅ Word document saved to: {output_path}")
